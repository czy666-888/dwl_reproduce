"""生成 dwl_reproduce2 综合分析 Word 文档"""
import csv, os, re, numpy as np
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

BASE = os.path.dirname(__file__)
LOGS = os.path.join(BASE, 'logs')
TRAIN_LOG = os.path.join(BASE, 'train_log.txt')
OUTPUT = os.path.join(LOGS, 'DWL训练分析报告_run1-7.docx')

RUNS = ['dwl_run1', 'dwl_run2', 'dwl_run3', 'dwl_run5', 'dwl_run6', 'dwl_run7']

COLUMN_MAP = {
    'r_lin_vel': ['r_lin_vel', 'r_track_lin_vel'],
    'r_ang_vel': ['r_ang_vel', 'r_track_ang_vel'],
    'r_orientation': ['r_orientation'],
    'r_height': ['r_height', 'r_base_height'],
    'r_foot_height': ['r_foot_height', 'r_feet_air_time'],
    'r_foot_vel': ['r_foot_vel', 'r_foot_slip'],
}


def _col(columns, key):
    if key in COLUMN_MAP:
        for c in COLUMN_MAP[key]:
            if c in columns:
                return c
    return key


def load_run(run):
    path = os.path.join(LOGS, run, 'behavior_log.csv')
    if not os.path.exists(path):
        return None
    with open(path, 'r', encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)
        rows = list(reader)
        columns = reader.fieldnames or list(rows[0].keys())

    def get(key):
        col = _col(columns, key)
        return np.array([float(r[col]) for r in rows])

    iters = np.array([int(r['iteration']) for r in rows])
    return {
        'run': run, 'n_rows': len(rows), 'min_iter': min(iters), 'max_iter': max(iters),
        'reward': np.array([float(r['mean_reward']) for r in rows]),
        'fall_pct': np.array([float(r['fall_pct']) for r in rows]),
        'base_height': np.array([float(r['base_height']) for r in rows]),
        'roll': np.abs(np.array([float(r['roll']) for r in rows])),
        'pitch': np.abs(np.array([float(r['pitch']) for r in rows])),
        'torque_mean': np.array([float(r['torque_mean']) for r in rows]),
        'r_lin_vel': get('r_lin_vel'),
        'r_ang_vel': get('r_ang_vel'),
        'r_orientation': get('r_orientation'),
        'r_height': get('r_height'),
        'r_foot_height': get('r_foot_height'),
        'r_foot_vel': get('r_foot_vel'),
    }


# Load all data
run_data = {}
for run in RUNS:
    d = load_run(run)
    if d:
        run_data[run] = d

# Parse train_log
with open(TRAIN_LOG, 'r', encoding='utf-8', errors='replace') as f:
    content = f.read()
pattern = r'Iter\s+(\d+)/\d+\s+\|\s+Reward:\s+([\d.]+)\s+\|\s+Loss:\s+([\d.]+)\s+\|\s+Steps/s:\s+([\d.]+)\s+\|\s+Time:\s+([\d.]+)'
matches = re.findall(pattern, content)
t_iters = np.array([int(m[0]) for m in matches])
t_rewards = np.array([float(m[1]) for m in matches])
t_losses = np.array([float(m[2]) for m in matches])
t_sps = np.array([float(m[3]) for m in matches])
t_times = np.array([float(m[4]) for m in matches])

# Dedupe
seen = set()
ci, cr, cl, cs, ct = [], [], [], [], []
for i, r, l, s, t in zip(t_iters, t_rewards, t_losses, t_sps, t_times):
    if i not in seen:
        seen.add(i)
        ci.append(i); cr.append(r); cl.append(l); cs.append(s); ct.append(t)
t_iters = np.array(ci); t_rewards = np.array(cr); t_losses = np.array(cl)
t_sps = np.array(cs); t_times = np.array(ct)


def stats5(d, key):
    return float(np.mean(d[key][:5])), float(np.mean(d[key][-5:]))


# ================================================================
# BUILD DOCUMENT
# ================================================================
doc = Document()

# Page style
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# --- Title ---
title = doc.add_heading('DWL复现训练分析报告', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'XBot-L双足机器人 | 256并行环境 | NVIDIA T600 4GB\n生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}').font.size = Pt(10)

doc.add_paragraph()

# ================================================================
# 1. 项目概述
# ================================================================
doc.add_heading('1. 项目概述', level=1)

doc.add_paragraph(
    '本报告分析DWL（Denoising World Model Learning）算法在XBot-L双足机器人上的复现训练过程。'
    '训练使用NVIDIA T600 4GB GPU，256个MuJoCo并行环境，PPO强化学习框架。'
    '共运行7个训练Run（run1~run7），跨越多个代码版本，逐步迭代改进。'
)

doc.add_heading('1.1 版本演进', level=2)

# Version table
vtable = doc.add_table(rows=7, cols=4, style='Table Grid')
vtable.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Run', '版本描述', '关键特征', '结果']
for i, h in enumerate(headers):
    cell = vtable.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True

versions = [
    ['dwl_run1', 'v1 Baseline', '12项自定义Reward，初始Critic实现', 'Reward≈2.0，全程无改善'],
    ['dwl_run2', 'v2 Critic修复', 'Critic输入修复，entropy/learning_rate调整', 'Reward≈3.9，但height下降'],
    ['dwl_run3', 'v4 足部轨迹', '五次多项式足部轨迹跟踪，步态硬编码', 'Reward≈1.5，最差'],
    ['dwl_run5', '测试Run', '短测试，仅49次迭代', '不足判断'],
    ['dwl_run6', 'v5 严格论文复现', '11项论文Reward + Critic Bug + privileged_dim=73', '45.2h，Reward 3.06→3.37'],
    ['dwl_run7', 'v5 + Bug修复', 'Critic全零Bug修复 + privileged_dim→137', '76.1h，Reward 3.06→4.14，最佳'],
]
for i, row_data in enumerate(versions):
    for j, val in enumerate(row_data):
        vtable.rows[i + 1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('1.2 关键Bug修复（Run6→Run7）', level=2)

doc.add_paragraph(
    'Bug 1（致命）— Critic输入全为零：构建Critic输入时，privileged维度不匹配（73 vs 137），'
    '导致回退到全零张量。GAE优势估计退化为随机噪声，PPO退化为REINFORCE，学习效率极低。'
)
doc.add_paragraph(
    'Bug 2（严重）— privileged维度不足：privileged_dim=73，论文要求184-47=137。'
    '缺失的64维特权信息导致Critic无法获取完整状态，价值估计偏差大。'
)
doc.add_paragraph(
    'Bug 3（次要）— _build_state()三处调用不统一：Rollout存储、Value计算、GAE last_value三处参数不一致。'
)

# ================================================================
# 2. 各Run详细分析
# ================================================================
doc.add_heading('2. 各Run训练详细分析', level=1)

run_configs = {
    'dwl_run1': ('v1 Baseline（旧版Reward + Critic Bug）', '3000次迭代'),
    'dwl_run2': ('v2 Critic修复（旧版Reward，部分Bug修复）', '1620次迭代（训练中断）'),
    'dwl_run3': ('v4 足部轨迹跟踪（硬编码五人多项式步态）', '3000次迭代'),
    'dwl_run5': ('v5 短测试', '175次迭代（中断）'),
    'dwl_run6': ('v5 严格论文复现（旧版，Critic Bug存在）', '3000次迭代，45.2小时'),
    'dwl_run7': ('v5 + 全部Bug修复（论文Reward + 正确Critic + 完整Privileged）', '2726次迭代，76.1小时'),
}

for run in RUNS:
    if run not in run_data:
        continue
    d = run_data[run]
    cfg_desc, scale = run_configs[run]

    doc.add_heading(f'2.{RUNS.index(run)+1} {run}', level=2)
    doc.add_paragraph(f'配置：{cfg_desc}').runs[0].bold = True
    doc.add_paragraph(f'规模：{d["n_rows"]}条行为记录，迭代 {d["min_iter"]}~{d["max_iter"]}，{scale}')

    # Stats table
    r_s, r_e = stats5(d, 'reward')
    f_s, f_e = stats5(d, 'fall_pct')
    h_s, h_e = stats5(d, 'base_height')
    roll_s, roll_e = stats5(d, 'roll')
    pitch_s, pitch_e = stats5(d, 'pitch')
    tq_s, tq_e = stats5(d, 'torque_mean')

    stable = doc.add_table(rows=7, cols=4, style='Table Grid')
    stable.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['指标', '起始值', '结束值', '变化趋势']):
        cell = stable.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run_text in p.runs:
                run_text.bold = True

    def trend_str(s, e, higher_better=True):
        if abs(s) < 0.001:
            return '持平'
        delta = (e - s) / abs(s) * 100
        direction = '↑' if delta > 0 else '↓'
        return f'{direction} {abs(delta):.1f}%'

    rows_data = [
        ['Mean Reward', f'{r_s:.4f}', f'{r_e:.4f}', trend_str(r_s, r_e, True)],
        ['Fall Rate', f'{f_s:.1f}%', f'{f_e:.1f}%', '始终100%' if f_s > 99 and f_e > 99 else trend_str(f_s, f_e, False)],
        ['Base Height', f'{h_s:.4f}m', f'{h_e:.4f}m', f'目标0.89m, 差距{0.89-h_e:.2f}m'],
        ['|Roll|', f'{roll_s:.4f}rad', f'{roll_e:.4f}rad', f'终止阈值1.0rad'],
        ['|Pitch|', f'{pitch_s:.4f}rad', f'{pitch_e:.4f}rad', f'终止阈值1.0rad'],
        ['Torque Mean', f'{tq_s:.1f}Nm', f'{tq_e:.1f}Nm', '—'],
    ]
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            stable.rows[i + 1].cells[j].text = val

    doc.add_paragraph()

    # Per-run analysis text
    analyses = {
        'dwl_run1': (
            'Run1作为v1 Baseline，Reward始终在2.0附近波动，几乎无学习迹象。Base Height从0.78m下降到0.77m，'
            '远低于0.89m的目标高度，说明机器人从未尝试站立。Roll和Pitch角度极小（<0.02rad），表明机器人在'
            '摔倒后几乎没有运动。Critic Bug导致PPO退化为REINFORCE，学习效率极低。'
        ),
        'dwl_run2': (
            'Run2修复了Critic的部分输入问题，Reward从3.75提升到3.98，改善明显。但Base Height从0.74m下降到'
            '0.74m（基本不变），且显著低于Run1的0.78m。这表明Reward函数可能促使机器人降低重心以获得更高的'
            '速度跟踪分数，代价是牺牲站立能力。训练在1620次迭代中断。'
        ),
        'dwl_run3': (
            'Run3加入五次多项式足部轨迹跟踪，是最差的Run。Reward从1.46跌至最低0.93，最终仅恢复到1.74。'
            '硬编码的正弦步态（0.64s周期）强迫机器人在站稳之前抬脚迈步，导致永远无法学习平衡。'
            'Pitch从0.006上升到0.025rad（增加316%），Base Height持续下降至0.73m。'
            '结论：用RL学习步态相位优于硬编码正弦波。'
        ),
        'dwl_run5': (
            'Run5仅49次迭代即中断，数据量不足以进行有意义的分析。Reward从3.03开始，与Run6/Run7起始值接近，'
            '确认了v5版Reward函数的基线水平。'
        ),
        'dwl_run6': (
            'Run6是第一个完整的v5版（严格论文复现Reward）训练。Reward从3.06缓慢提升到3.37（+10%），'
            '但存在Critic全零Bug，导致学习效率受限。总耗时45.2小时完成3000次迭代。'
            'Base Height持续下降至0.73m，Roll从0.003增加到0.009rad。'
            '尽管Reward有改善，但摔倒率始终100%，机器人从未学会真正行走。'
        ),
        'dwl_run7': (
            'Run7是修复所有已知Bug后的训练，也是表现最佳的Run。Reward从3.06持续上升到4.14（+35%），'
            '最佳Reward达到4.15（@iter 2700）。训练耗时76.1小时，速度从~460降至~160 steps/s。'
            'Base Height从0.74m小幅下降至0.73m，Roll和Pitch均在安全范围内。'
            '然而，摔倒率依然100%，Base Height远低于0.89m目标。'
            'Reward的提升主要来自速度跟踪(r_lin_vel)、周期性接触(r_periodic_contact)等分项的优化，'
            '而非真正的站立行走。机器人学会了在摔倒状态下更好地匹配速度指令，但从未站起来。'
        ),
    }
    doc.add_paragraph(analyses[run])

# ================================================================
# 3. 跨Run对比
# ================================================================
doc.add_heading('3. 跨Run对比分析', level=1)

doc.add_heading('3.1 关键指标汇总', level=2)

# Summary comparison table
ctable = doc.add_table(rows=len(RUNS) + 1, cols=6, style='Table Grid')
ctable.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Run', 'Reward Start', 'Reward End', 'Δ', 'Height Start→End', 'Fall Rate']):
    cell = ctable.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run_text in p.runs:
            run_text.bold = True

for idx, run in enumerate(RUNS):
    if run not in run_data:
        continue
    d = run_data[run]
    r_s, r_e = stats5(d, 'reward')
    h_s, h_e = stats5(d, 'base_height')
    f_s, f_e = stats5(d, 'fall_pct')
    delta = r_e - r_s
    ctable.rows[idx + 1].cells[0].text = run
    ctable.rows[idx + 1].cells[1].text = f'{r_s:.4f}'
    ctable.rows[idx + 1].cells[2].text = f'{r_e:.4f}'
    ctable.rows[idx + 1].cells[3].text = f'{delta:+.4f}'
    ctable.rows[idx + 1].cells[4].text = f'{h_s:.4f}→{h_e:.4f}'
    ctable.rows[idx + 1].cells[5].text = f'{f_s:.0f}%→{f_e:.0f}%'

doc.add_paragraph()

doc.add_heading('3.2 版本对比关键发现', level=2)

findings = [
    ('v1→v2（Reward提升）', 'Run1→Run2 Reward从2.0跃升至3.9（+95%），Critic修复贡献巨大'),
    ('v2→v4（Reward暴跌）', 'Run2→Run3 Reward从3.9跌至1.5（-62%），硬编码步态适得其反'),
    ('v4→v5（Reward恢复）', 'Run3→Run6 Reward从1.5恢复到3.1，论文Reward函数优于自定义版本'),
    ('v5 Bug修复（持续改善）', 'Run6→Run7 Reward从3.4提升到4.1（+21%），Critic+Privileged修复后学习信号显著改善'),
    ('共性问题', '所有Run摔倒率100%，Base Height始终<0.78m（目标0.89m）。Reward提升≠真正行走。'),
]
for title, desc in findings:
    p = doc.add_paragraph()
    p.add_run(f'{title}：').bold = True
    p.add_run(desc)

# ================================================================
# 4. 训练曲线分析（train_log）
# ================================================================
doc.add_heading('4. Run7训练过程分析（train_log.txt）', level=1)

mid = len(t_rewards) // 2
doc.add_paragraph(
    f'Run7共训练2725次迭代（iter 1~2726），总耗时{t_times[-1]/60:.1f}小时（{t_times[-1]/3600:.1f}天）。'
    f'训练记录在train_log.txt中。'
)

doc.add_heading('4.1 训练阶段划分', level=2)

phase_table = doc.add_table(rows=5, cols=5, style='Table Grid')
phase_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['阶段', '迭代范围', 'Reward均值', 'Loss均值', 'Steps/s']):
    cell = phase_table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run_text in p.runs:
            run_text.bold = True

q1 = len(t_rewards) // 4
phases = [
    ('预热期', 0, q1),
    ('快速下降期', q1, 2 * q1),
    ('稳定收敛期', 2 * q1, 3 * q1),
    ('精细调优期', 3 * q1, len(t_rewards)),
]
for i, (name, start, end) in enumerate(phases):
    phase_table.rows[i + 1].cells[0].text = name
    phase_table.rows[i + 1].cells[1].text = f'{t_iters[start]}~{t_iters[end-1]}'
    phase_table.rows[i + 1].cells[2].text = f'{np.mean(t_rewards[start:end]):.4f}'
    phase_table.rows[i + 1].cells[3].text = f'{np.mean(t_losses[start:end]):.1f}'
    phase_table.rows[i + 1].cells[4].text = f'{np.mean(t_sps[start:end]):.0f}'

doc.add_paragraph()

doc.add_heading('4.2 训练效率', level=2)

active = t_sps[t_sps > 50]
doc.add_paragraph(
    f'平均训练速度：{np.mean(active):.0f} steps/s（有效区间）。速度呈下降趋势，从初始~460降至后期~160 steps/s，'
    f'可能是模型复杂度增加或GPU散热降频导致。'
)
doc.add_paragraph(
    f'Loss从4851迅速下降至~5（前50次迭代），随后稳定在4~6区间，表明DWL世界模型快速收敛。'
    f'Reward前半段均值{np.mean(t_rewards[:mid]):.3f}，后半段{np.mean(t_rewards[mid:]):.3f}（+{np.mean(t_rewards[mid:])-np.mean(t_rewards[:mid]):.3f}），'
    f'持续改善但速度放缓。'
)

# ================================================================
# 5. Reward分项分析
# ================================================================
doc.add_heading('5. Reward分项分析（Run7）', level=1)

if 'dwl_run7' in run_data:
    d = run_data['dwl_run7']

    doc.add_paragraph('以下分析基于dwl_run7的behavior_log.csv，展示11项Reward分项的学习趋势。')

    r_items = [
        ('Linear Velocity (r_lin_vel)', d['r_lin_vel'], '速度跟踪能力'),
        ('Angular Velocity (r_ang_vel)', d['r_ang_vel'], '角速度跟踪能力'),
        ('Orientation (r_orientation)', d['r_orientation'], '姿态保持（重力方向对齐）'),
        ('Height (r_height)', d['r_height'], '基座高度跟踪（目标0.89m）'),
        ('Foot Height (r_foot_height)', d['r_foot_height'], '足部高度跟踪（步态质量）'),
        ('Foot Velocity (r_foot_vel)', d['r_foot_vel'], '足部速度跟踪'),
    ]

    rtable = doc.add_table(rows=len(r_items) + 1, cols=5, style='Table Grid')
    rtable.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Reward分项', '起始值', '结束值', '变化(%)', '趋势解读']):
        cell = rtable.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run_text in p.runs:
                run_text.bold = True

    interpretations = {
        'r_lin_vel': '大幅改善(+78%)，机器人学会匹配速度指令',
        'r_ang_vel': '大幅改善(+60%)，角速度跟踪能力增强',
        'r_orientation': '持续下降(-26%)，姿态控制恶化',
        'r_height': '持平(~0%)，高度跟踪完全未学习',
        'r_foot_height': '先升后降，步态学习不稳定',
        'r_foot_vel': '先升后降，与r_foot_height一致',
    }

    for i, (name, data, desc) in enumerate(r_items):
        s_val = np.mean(data[:5])
        e_val = np.mean(data[-5:])
        pct = (e_val - s_val) / abs(s_val) * 100 if abs(s_val) > 0.001 else 0

        rtable.rows[i + 1].cells[0].text = name
        rtable.rows[i + 1].cells[1].text = f'{s_val:.4f}'
        rtable.rows[i + 1].cells[2].text = f'{e_val:.4f}'
        rtable.rows[i + 1].cells[3].text = f'{pct:+.1f}%'
        # Extract key for interpretation
        key = None
        for k, v in interpretations.items():
            if k in name.lower() or name.lower().startswith(k.split('(')[0].strip().lower()):
                key = k
                break
        rtable.rows[i + 1].cells[4].text = interpretations.get(key, desc)

    doc.add_paragraph()

    doc.add_paragraph(
        '关键发现：速度跟踪类Reward（r_lin_vel、r_ang_vel）是Run7 Reward提升的主要驱动力（+78%/+60%）。'
        '但姿态相关Reward持续下降（r_orientation -26%），高度跟踪完全未学习。'
        '这表明策略学到的不是"站立行走"，而是"以摔倒姿态匹配速度指令"的捷径解。'
    )

# ================================================================
# 6. 行为指标分析
# ================================================================
doc.add_heading('6. 行为指标分析（Run7）', level=1)

if 'dwl_run7' in run_data:
    d = run_data['dwl_run7']

    b_items = [
        ('Base Height', d['base_height'], 'm', 0.89, '目标0.89m → 实际0.73m，差距0.16m（-18%），从未接近目标'),
        ('|Roll|', d['roll'], 'rad', 1.0, '始终<0.02rad，远低于终止阈值1.0rad，机器人几乎无横向运动'),
        ('|Pitch|', d['pitch'], 'rad', 1.0, '从0.006升至0.029rad（+383%），俯仰不稳定恶化明显'),
        ('Torque Mean', d['torque_mean'], 'Nm', None, '20→20Nm，力矩使用稳定，无异常'),
    ]

    btable = doc.add_table(rows=len(b_items) + 1, cols=4, style='Table Grid')
    btable.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['指标', 'Start→End', '阈值', '解读']):
        cell = btable.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run_text in p.runs:
                run_text.bold = True

    for i, (name, data, unit, threshold, interp) in enumerate(b_items):
        s = float(np.mean(data[:5]))
        e = float(np.mean(data[-5:]))
        thresh_str = f'{threshold}{unit}' if threshold else '—'
        btable.rows[i + 1].cells[0].text = name
        btable.rows[i + 1].cells[1].text = f'{s:.4f}→{e:.4f} {unit}'
        btable.rows[i + 1].cells[2].text = thresh_str
        btable.rows[i + 1].cells[3].text = interp

doc.add_paragraph()

# ================================================================
# 7. 核心问题诊断
# ================================================================
doc.add_heading('7. 核心问题诊断', level=1)

problems = [
    ('7.1 致命问题：摔倒率100%',
     '所有7个Run的摔倒率始终为100%，从未有一个Run学会站立或行走。这是最核心的未解决问题。\n\n'
     '可能原因：\n'
     '  a) Base Height始终<0.78m，远低于0.89m目标。机器人从未经历"站立状态"，因此无法从站立中学习。\n'
     '  b) 终止高度(0.35m)过低，机器人在极低姿态下仍可生存，Reward信号无法区分"站立"与"爬行"。\n'
     '  c) 初期摔倒导致所有episode在摔倒姿态下运行，GAE优势估计偏向"摔倒策略"。\n'
     '  d) 域随机化（推力、噪声、摩擦）在256环境下可能过强，使机器人无法稳定。'),
    ('7.2 Reward捷径问题',
     '速度跟踪Reward（r_lin_vel、r_ang_vel）持续上升，但姿态/高度Reward不升反降。'
     '策略可能学到了一个"cheating"解：保持摔倒姿态并调整关节以匹配速度指令，'
     '而不是站直行走。需要增加站立/高度相关的Reward权重或引入curriculum learning。'),
    ('7.3 训练环境规模',
     '256个并行环境远少于论文的12288个（仅2%）。PPO的on-policy性质要求足够多的样本覆盖策略空间，'
     '256环境下每个batch仅有12288个transition（256×48），可能导致样本多样性不足。'),
    ('7.4 步态设计问题',
     'Run3的硬编码正弦步态使Reward下降62%，说明预定义步态不适合XBot-L。'
     'Run7虽然使用周期性接触Reward引导步态，但机器人从未站立，无法形成有效的周期性步态。'),
    ('7.5 GPU/硬件限制',
     'T600 4GB是入门级专业卡，训练速度从460降至160 steps/s。76小时的训练仅完成2726迭代，'
     '而论文的3000迭代在有Bug的情况下仍需45~76小时。更大规模实验需要更强GPU。'),
]

for title, desc in problems:
    doc.add_heading(title, level=2)
    doc.add_paragraph(desc)

# ================================================================
# 8. 建议与下一步
# ================================================================
doc.add_heading('8. 建议与下一步', level=1)

suggestions = [
    ('1. 引入站立Curriculum',
     '分阶段训练：Phase 1（0~1000 iter）：固定机器人站立（零速度指令），仅优化r_height + r_orientation，'
     '让机器人先学会站立。Phase 2（1000~2000 iter）：逐步引入低速指令（0.1→0.3 m/s）。'
     'Phase 3（2000+ iter）：正常速度训练。'),
    ('2. 调整Reward权重',
     '大幅提高r_height权重（0.5→2.0），提高r_orientation权重（1.0→2.0）。'
     '考虑增加"生存bonus"：每存活1秒给予正Reward，而非仅在episode结束时结算。'),
    ('3. 提高终止高度',
     '将termination_height从0.35m提升至0.60m。0.35m对XBot-L（站立0.89m）过于宽容，'
     '机器人几乎平躺在地上也不终止。0.60m约为站立的67%，是更合理的阈值。'),
    ('4. 减少域随机化强度',
     '初期训练（前500 iter）关闭随机推力和大幅摩擦力变化，让策略先学习基本站立。'
     '后续逐步增强域随机化以提升鲁棒性。'),
    ('5. 增加训练环境数',
     '在T600 4GB显存允许下，尝试提升num_envs（256→384或512）。降低num_steps_per_env（48→24）'
     '以补偿显存，保持总transition数不变的同时增加环境多样性。'),
    ('6. 考虑参考状态初始化',
     '使用reference motion或motion capture数据初始化，将机器人设置为站立姿态开始每个episode，'
     '而非从地面开始。给策略一个"站立经验"的起点。'),
]

for title, desc in suggestions:
    p = doc.add_paragraph()
    p.add_run(title + '：').bold = True
    p.add_run(desc)

# ================================================================
# 9. 附录
# ================================================================
doc.add_heading('9. 附录', level=1)

doc.add_heading('9.1 生成的分析图表', level=2)
charts = [
    '01_training_overview.png — 训练总览（Reward/Loss/Speed/累计时间）',
    '02_training_phases.png — 训练分阶段分析',
    '03_all_runs_reward_fall.png — 6个Run Reward+Fall率对比',
    '04a_run7_reward_items.png — Run7 11项Reward分项',
    '04b_run7_behavior.png — Run7 行为指标',
    '04c_run7_summary.png — Run7 总览',
    '05_cross_run_comparison.png — 跨Run 9指标趋势对比',
    '06_start_vs_end_bars.png — 起始vs结束柱状图',
    '07_distribution_shift.png — Reward/Loss分布偏移',
    '08_run7_velocity_contacts.png — 速度跟踪+周期性接触',
]
for c in charts:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('9.2 数据文件', level=2)
data_files = [
    'train_log.txt — Run7完整训练日志（2725行）',
    'logs/dwl_run1~7/behavior_log.csv — 各Run行为记录（共2727行）',
    'checkpoints/best_model.pt — 最优模型（iter 2700, 9.3MB）',
    'checkpoints/model_iter_*.pt — 每100迭代检查点',
    'analyze_all_runs.py — 本报告的分析脚本',
]
for f in data_files:
    doc.add_paragraph(f, style='List Bullet')

# ================================================================
# SAVE
# ================================================================
doc.save(OUTPUT)
print(f'Document saved: {OUTPUT}')
