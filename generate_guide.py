"""生成 DWL 代码入门指南 Word 文档"""
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ============================================================
# TITLE
# ============================================================
title = doc.add_heading('DWL复现项目 — 代码入门指南', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(f'XBot-L双足机器人 | PPO强化学习 | 论文复现\n生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}').font.size = Pt(10)

doc.add_paragraph(
    '本文档是面向代码小白的入门级解释，涵盖：项目整体架构、PPO训练流程、Action（动作）、'
    'Reward（奖励函数）的11项构成、以及训练规模详解。配合之前的《DWL训练分析报告_run1-7.docx》一起阅读。'
)

# ============================================================
# 1. 项目在做什么
# ============================================================
doc.add_heading('1. 你在做什么？', level=1)
doc.add_paragraph(
    '你在复现一篇学术论文，让一个虚拟的双足机器人在物理模拟器里学会走路。'
    '你不是手动编程告诉它每一步怎么迈——而是让它自己尝试、摔倒、再尝试，通过几千万次试错自己"悟"出怎么走路。'
    '这就是强化学习（Reinforcement Learning）。'
)
doc.add_paragraph(
    '用教狗"坐下"来类比：你不是手把手摆狗的身体，而是在它做对时给零食（Reward），做错时不給。'
    '狗（AI）自己会慢慢摸索出"坐着就有零食吃"这件事。你的机器人也一样——'
    '做对动作就给高分，做错就给低分，它自己学着追求高分。'
)

# ============================================================
# 2. 代码四大模块
# ============================================================
doc.add_heading('2. 代码怎么工作的？（四个模块）', level=1)

doc.add_heading('2.1 环境 (dwl_env.py) — "训练场"', level=2)
doc.add_paragraph(
    '创建256个机器人"副本"同时训练（等于256只狗同时学）。使用MuJoCo物理引擎模拟真实的重力、碰撞、摩擦力。'
    '每0.01秒（100Hz）告诉机器人12个关节怎么动。给机器人一个速度指令（如"往前0.3m/s"），看它能不能跟得上。'
    '如果机器人摔倒（高度 < 0.35m 或倾斜超过57度），这一回合就结束。'
)

doc.add_heading('2.2 大脑 (dwl_networks.py) — "机器人的神经网络"', level=2)
doc.add_paragraph('这是机器人的"大脑"，由4个子网络组成：')
# Table
btable = doc.add_table(rows=5, cols=3, style='Table Grid')
btable.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['子网络', '输入 → 输出', '类比']):
    btable.rows[0].cells[i].text = h
    for p in btable.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
brain_data = [
    ['Encoder (编码器)', '47个观测数字 → 24个理解信息', '狗的眼睛+大脑：看到周围环境，理解发生了什么'],
    ['Actor (策略网络)', '24个理解信息 → 12个动作数字', '狗的决策：决定现在做什么'],
    ['Critic (价值网络)', '184个完整状态 → 1个分数', '狗的直觉：预感当前状态好不好'],
    ['Decoder (解码器)', '24个理解信息 → 184个世界重建', '想象"我做这个动作会带来什么结果"'],
]
for i, row in enumerate(brain_data):
    for j, val in enumerate(row):
        btable.rows[i+1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('2.3 PPO训练器 (dwl_ppo.py) — "教练"', level=2)
doc.add_paragraph(
    'PPO = Proximal Policy Optimization（近端策略优化），是当前最主流的强化学习算法之一。'
)
p1 = doc.add_paragraph()
p1.add_run('核心思想：').bold = True
p1.add_run(
    '鼓励好行为（高分动作），但控制变化速度（不要一夜之间大变样）。'
    '就像教练对学生说"这个动作不错，多练练"，但也说"别一下子改太多，慢慢来"。'
)
p2 = doc.add_paragraph()
p2.add_run('训练流程（每次Iteration）：').bold = True
doc.add_paragraph('1. 采集阶段：256个机器人各跑48步，记录所有数据（观测、动作、奖励、是否摔倒）', style='List Number')
doc.add_paragraph('2. 评分阶段：GAE算法回顾每一步——"你这个动作是真的好，还是恰好运气好？"', style='List Number')
doc.add_paragraph('3. 学习阶段：神经网络更新参数，提高高分动作的出现概率', style='List Number')
doc.add_paragraph('4. 重复3000次', style='List Number')

p3 = doc.add_paragraph()
p3.add_run('GAE（广义优势估计）：').bold = True
p3.add_run(
    '用来区分"能力强"和"运气好"。如果你在机器人快摔倒时做了一个好动作挽救了局面，这个动作得分就高。'
    '如果你只是在一个本来就很安全的状态下随便动一下，得分就低。'
)
p4 = doc.add_paragraph()
p4.add_run('PPO的Clip机制：').bold = True
p4.add_run(
    '限制策略改变的幅度。比如原来的策略是"50%概率抬左脚"，PPO会限制你最多改到70%或30%，'
    '防止一夜之间策略大变样导致学歪。'
)

doc.add_heading('2.4 训练主循环 (train.py) — "总指挥"', level=2)
doc.add_paragraph('把上面三部分串起来的控制脚本：')
doc.add_paragraph('1. 初始化：创建256个环境 + 神经网络 + 训练器', style='List Number')
doc.add_paragraph('2. 主循环（3000次）：', style='List Number')
doc.add_paragraph('   a. 256个机器人各跑48步 → 采集12,288条经验', style='List')
doc.add_paragraph('   b. GAE计算每条经验的"真实分数"', style='List')
doc.add_paragraph('   c. 神经网络学习（2轮 × 4批 = 8次参数更新）', style='List')
doc.add_paragraph('   d. 每5次Iteration存一次行为数据到CSV', style='List')
doc.add_paragraph('   e. 每100次Iteration保存一次模型', style='List')
doc.add_paragraph('   f. 打印进度到屏幕和train_log.txt', style='List')
doc.add_paragraph('3. 训练结束：保存最终模型和最优模型', style='List Number')

# ============================================================
# 3. Action
# ============================================================
doc.add_heading('3. Action（动作）是什么？', level=1)

doc.add_paragraph(
    'Action就是12个数字，控制机器人12条腿部关节的转动角度。每一步（0.01秒），AI输出一个长度为12的向量。'
)

doc.add_paragraph('XBot-L的12个关节（每条腿6个）：')
atable = doc.add_table(rows=7, cols=3, style='Table Grid')
atable.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['关节名', '位置', '作用']):
    atable.rows[0].cells[i].text = h
    for p in atable.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
joints = [
    ['leg_roll', '髋部侧向', '腿左右摆'],
    ['leg_yaw', '髋部旋转', '腿内外旋'],
    ['leg_pitch', '髋部前后', '腿前后摆（走路主驱动）'],
    ['knee', '膝盖', '弯腿/伸直'],
    ['ankle_pitch', '脚踝前后', '脚掌上下'],
    ['ankle_roll', '脚踝侧向', '脚掌左右'],
]
for i, row in enumerate(joints):
    atable.rows[i+1].cells[0].text = row[0]
    atable.rows[i+1].cells[1].text = row[1]
    atable.rows[i+1].cells[2].text = row[2]

doc.add_paragraph()

doc.add_paragraph('一个Action的示例：')
doc.add_paragraph('[0.12, -0.05, 0.08, -0.30, 0.02, 0.01,  -0.10, 0.04, 0.07, -0.28, 0.01, -0.02]')
doc.add_paragraph('  ← 左腿6个关节 →                          ← 右腿6个关节 →')

doc.add_paragraph()
doc.add_paragraph('Action是怎么变成物理动作的：')
doc.add_paragraph('AI输出(action) → 乘以缩放因子(0.25) → 加上默认姿态(站直) → 得到目标角度 → PD控制器计算力矩 → 电机施加力矩 → MuJoCo物理模拟', style='List')

# ============================================================
# 4. Reward
# ============================================================
doc.add_heading('4. Reward（奖励函数）怎么设计？', level=1)

doc.add_paragraph(
    '奖励函数就是给机器人的"打分标准"。每一步动作做完后，从11个方面评分，加权求和得到总分。'
    '核心设计思想来自论文 Table V。每一项都使用指数函数 exp(-误差² × sigma)，让分数永远在0~1之间——'
    '越接近1越好，越接近0越差。'
)

# Reward table
rtable = doc.add_table(rows=12, cols=5, style='Table Grid')
rtable.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['#', '奖励项', '权重', '问什么？', '公式核心']):
    rtable.rows[0].cells[i].text = h
    for p in rtable.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True

reward_rows = [
    ['1', '线速度跟踪\n(r_lin_vel)', '1.0',
     '实际速度跟指令差多少？\n例：指令0.3m/s，实际0.3 → 满分',
     'exp(-误差² × 5.0)'],
    ['2', '角速度跟踪\n(r_ang_vel)', '1.0',
     '转向速度跟指令差多少？',
     'exp(-误差² × 7.0)'],
    ['3', '姿态保持\n(r_orientation)', '1.0',
     '身体有没有歪？\nroll/pitch是否接近0？',
     'exp(-(roll²+pitch²) × 5.0)'],
    ['4', '身高跟踪\n(r_height)', '0.5',
     '站得够不够高？\n目标0.89m，实际0.73m',
     'exp(-(高度差)² × 10.0)'],
    ['5', '周期性接触力\n(r_periodic_contact)', '1.0',
     '支撑脚有没有踩实地面？\n走路时双脚交替支撑',
     'stance × 接触力'],
    ['6', '周期性足部速度\n(r_periodic_vel)', '1.0',
     '抬起的脚有没有在移动？\n迈步时脚应该有速度',
     '(1-stance) × 足部速度'],
    ['7', '足部高度跟踪\n(r_foot_height)', '1.0',
     '抬脚高度是否跟上参考轨迹？\n五次多项式曲线，最高6cm',
     'exp(-高度误差² × 5.0)'],
    ['8', '足部速度跟踪\n(r_foot_vel)', '0.5',
     '脚抬起/落下的速度是否合适？',
     'exp(-速度误差² × 3.0)'],
    ['9', '默认关节姿态\n(r_default_joint)', '0.2',
     '关节有没有乱动？\n保持接近站直姿态',
     'exp(-关节差² × 2.0)'],
    ['10', '能耗惩罚\n(r_energy)', '-0.0001',
     '用了多少力气？\n唯一带负号权重的项',
     '|力矩 × 转速| × (-1)'],
    ['11', '动作平滑\n(r_action_smooth)', '-0.01',
     '相邻两次动作变化大不大？\n防止抽搐',
     '动作加速度² × (-1)'],
]
for i, row in enumerate(reward_rows):
    for j, val in enumerate(row):
        rtable.rows[i+1].cells[j].text = val

doc.add_paragraph()
doc.add_paragraph('总Reward公式：')
doc.add_paragraph(
    'Total = 1.0×线速度 + 1.0×角速度 + 1.0×姿态 + 0.5×身高\n'
    '      + 1.0×接触力 + 1.0×足部速度 + 1.0×足部高度 + 0.5×足速度跟踪\n'
    '      + 0.2×默认姿态 - 0.0001×能耗 - 0.01×动作平滑'
)

doc.add_paragraph()
p5 = doc.add_paragraph()
p5.add_run('关键发现（来自Run7数据）：').bold = True
doc.add_paragraph('速度跟踪类Reward是主要驱动力：r_lin_vel +78%, r_ang_vel +60%', style='List Bullet')
doc.add_paragraph('姿态类Reward持续下降：r_orientation -26%', style='List Bullet')
doc.add_paragraph('身高类Reward完全未学习：r_height ~0%变化', style='List Bullet')
doc.add_paragraph('结论：机器人学到的是"以摔倒姿态匹配速度指令"的捷径解，而非真正的站立行走', style='List Bullet')

# ============================================================
# 5. 训练规模
# ============================================================
doc.add_heading('5. 训练规模详解', level=1)

doc.add_heading('5.1 三层"步"的概念（最容易搞混）', level=2)
stable = doc.add_table(rows=4, cols=4, style='Table Grid')
stable.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['层级', '频率', '间隔', '说明']):
    stable.rows[0].cells[i].text = h
    for p in stable.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
steps_data = [
    ['物理步 (sim_dt)', '1000Hz', '0.001秒', 'MuJoCo引擎每一步计算重力、碰撞、摩擦力'],
    ['策略步 (policy_dt)', '100Hz', '0.01秒', 'AI每10个物理步做一次决策（decimation=10）'],
    ['训练步 (Iteration)', '~0.3次/分钟', '3~5分钟', '每48个策略步学习一次'],
]
for i, row in enumerate(steps_data):
    for j, val in enumerate(row):
        stable.rows[i+1].cells[j].text = val

doc.add_paragraph()

doc.add_heading('5.2 一次Iteration（一轮训练）的数据量', level=2)
doc.add_paragraph('一个Iteration的流程：')
doc.add_paragraph('256个机器人 × 每个跑48个策略步 = 共12,288次AI决策', style='List')
doc.add_paragraph('12,288 × 10次物理步 = 122,880步物理模拟', style='List')
doc.add_paragraph('模拟了 12,288 × 0.01s = 122.88秒 的机器人体验', style='List')
doc.add_paragraph('收集12,288条经验 → GAE评分 → 神经网络学习（2轮 × 4批 = 8次更新）', style='List')
doc.add_paragraph('真实耗时：约3-5分钟（取决于GPU速度）', style='List')

doc.add_heading('5.3 完整训练（3000次Iteration）', level=2)

scale_table = doc.add_table(rows=6, cols=2, style='Table Grid')
scale_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['指标', '数值']):
    scale_table.rows[0].cells[i].text = h
    for p in scale_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
scale_data = [
    ['总AI决策次数', '3000 × 12,288 = 36,864,000 次（约3686万次）'],
    ['总物理模拟步数', '3000 × 122,880 = 368,640,000 步（约3.7亿步）'],
    ['总模拟时间', '36,864,000 × 0.01s ≈ 102.4 小时（机器人"生活"了4天多）'],
    ['真实GPU时间', '约76小时（T600 4GB笔记本显卡）'],
    ['学习次数', '3000 × 8 = 24,000 次神经网络参数更新'],
]
for i, row in enumerate(scale_data):
    scale_table.rows[i+1].cells[0].text = row[0]
    scale_table.rows[i+1].cells[1].text = row[1]

doc.add_paragraph()

doc.add_heading('5.4 一个Episode（一个机器人的"一辈子"）', level=2)
doc.add_paragraph(
    '一个Episode = 一个机器人从出生到结束的完整过程：最长2400个策略步（24秒模拟时间）。'
    '但由于摔倒率100%，实际情况是：'
)
doc.add_paragraph('机器人出生在0.95m高度', style='List Bullet')
doc.add_paragraph('0.2~0.5秒内摔倒（高度 < 0.35m触发终止）', style='List Bullet')
doc.add_paragraph('剩余23秒多趴在地上，Reward全为0', style='List Bullet')
doc.add_paragraph('这就是为什么Base Height永远只有0.73m——机器人从未体验过站立的感受', style='List Bullet')

doc.add_heading('5.5 论文 vs 你的配置差距', level=2)
comp_table = doc.add_table(rows=4, cols=4, style='Table Grid')
comp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['', '论文', '你的', '比例']):
    comp_table.rows[0].cells[i].text = h
    for p in comp_table.rows[0].cells[i].paragraphs:
        for r in p.runs:
            r.bold = True
comp_data = [
    ['并行环境数', '12,288', '256', '2%'],
    ['每轮采集经验', '589,824', '12,288', '2%'],
    ['GPU', '数据中心级', 'T600 4GB（笔记本）', '差距巨大'],
]
for i, row in enumerate(comp_data):
    for j, val in enumerate(row):
        comp_table.rows[i+1].cells[j].text = val
doc.add_paragraph()
doc.add_paragraph(
    '论文一次Iteration的样本量是你的48倍。PPO的on-policy性质要求足够多样本覆盖策略空间。'
    '256个环境的样本多样性不足，机器人碰不到足够多"差点摔倒又救回来"的边缘情况，难以学到稳定的平衡策略。'
)

# ============================================================
# 6. 完整流程图
# ============================================================
doc.add_heading('6. 完整训练流程（一图看懂）', level=1)

doc.add_paragraph('每一步（0.01秒）：')
doc.add_paragraph('① AI看到47个数字（关节角度、速度、姿态、指令等）→ 观测', style='List Number')
doc.add_paragraph('② Encoder(GRU)处理历史观测 → 输出24维隐状态z', style='List Number')
doc.add_paragraph('③ Actor根据z → 输出12个动作数字', style='List Number')
doc.add_paragraph('④ 12个数字 × 0.25 + 默认姿态 → 12个目标关节角度', style='List Number')
doc.add_paragraph('⑤ PD控制器 → 12个关节力矩 → MuJoCo物理模拟10步(0.001s×10)', style='List Number')
doc.add_paragraph('⑥ 计算11项Reward → 加权求和 = 一个总分', style='List Number')
doc.add_paragraph('⑦ 检查是否摔倒（高度<0.35m 或 tilt>57°）', style='List Number')
doc.add_paragraph('⑧ 回到①（重复48步后触发一次学习）', style='List Number')

doc.add_paragraph()
doc.add_paragraph('每48步后（一次Iteration）：')
doc.add_paragraph('⑨ 12,288条经验 → GAE计算优势函数 → 2轮×4批更新参数', style='List Number')

doc.add_paragraph()
doc.add_paragraph('每3000次Iteration后：')
doc.add_paragraph('⑩ 保存模型，训练结束', style='List Number')

# ============================================================
# 7. 当前问题
# ============================================================
doc.add_heading('7. 当前核心问题总结', level=1)

doc.add_paragraph(
    '尽管Run7的Reward持续上升（3.06→4.14），但存在根本性问题：'
)

problems_summary = [
    ('摔倒率100%',
     '所有Run从未有一个机器人学会站立。终止高度0.35m过低（站立高度0.89m的39%），机器人几乎平躺也能存活。'),
    ('Reward捷径',
     '机器人学会了"趴着匹配速度指令"——速度跟踪Reward +78%，但姿态Reward -26%，身高Reward无改善。'
     '策略找到了一条"作弊"路径，而非真正的站立行走。'),
    ('样本量不足',
     '256个并行环境仅为论文的2%。PPO依赖大量独立样本，256环境下的策略多样性不足以覆盖双足行走的状态空间。'),
    ('从未体验站立',
     'Base Height始终<0.78m。机器人出生后很快摔倒，从来没经历过"站着的状态"，也就无法从中学到平衡。'),
    ('步态系统矛盾',
     '硬编码的0.64s步态周期强迫抬脚（Run3验证了其失败），但Run7的周期性接触Reward也未能引导出自然步态——'
     '因为机器人从未站立，步态无从谈起。'),
]
for title, desc in problems_summary:
    p = doc.add_paragraph()
    p.add_run(title + '：').bold = True
    p.add_run(desc)

# ============================================================
# SAVE
# ============================================================
output = r'C:\Users\czy66\Desktop\dwl_reproduce2\DWL代码入门指南.docx'
doc.save(output)
print(f'Saved: {output}')
