from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os, datetime

desktop = os.environ['USERPROFILE'] + r'\Desktop'
doc = Document()

# Style setup
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(10.5)

# ==== Title ====
title = doc.add_heading('DWL论文复现总结报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph('生成日期：2026年6月7日')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p = doc.add_paragraph('基于MuJoCo的Denoising World Model Learning复现方案')
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# ==== Chapter 1: Paper Overview ====
doc.add_heading('一、论文基本信息', 1)
info_data = [
    ('论文标题',
     'Advancing Humanoid Locomotion: Mastering Challenging Terrains with Denoising World Model Learning'),
    ('发表会议', 'Robotics: Science and Systems (RSS) 2024'),
    ('作者单位', '清华大学 / 上海期智研究院 / RobotEra(星动纪元)'),
    ('实验平台', 'XBot-S (1.2m/38kg) & XBot-L (1.65m/57kg)'),
    ('核心思路', '全栈端到端强化学习 + 零样本Sim-to-Real迁移'),
    ('算法框架', 'DWL = GRU Encoder(去噪+压缩) + Decoder(状态重建) + PPO(Actor/Critic)'),
    ('关键创新', 'Denoising World Model：将仿真中的噪声观测去噪压缩为24维隐状态，结合PPO联合训练'),
    ('官方代码', 'https://github.com/roboterax/humanoid-gym (RobotEra官方)'),
]
table = doc.add_table(rows=len(info_data), cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate(info_data):
    table.rows[i].cells[0].text = k
    table.rows[i].cells[1].text = v
    for cell in [table.rows[i].cells[0]]:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

# ==== Chapter 2: Original Requirements ====
doc.add_heading('二、原版复现必要条件', 1)

doc.add_heading('2.1 硬件要求', 2)
hw_items = [
    'GPU：NVIDIA GPU（推荐A100/4090），论文使用12288并行环境需大显存（24GB+）',
    'CPU：多核CPU（至少10线程用于PhysX物理计算）',
    '内存：32GB+（12288环境状态缓冲）',
    '存储：50GB+（Isaac Gym + 模型训练checkpoint）',
]
for item in hw_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.2 软件要求', 2)
sw_items = [
    '操作系统：Ubuntu 20.04/22.04（Isaac Gym仅支持Linux）',
    'CUDA Toolkit：11.x 或 12.x',
    '仿真引擎：Isaac Gym Preview 4（NVIDIA私有软件，需申请开发者计划）',
    'PyTorch：GPU版（CUDA-enabled），需与Isaac Gym版本兼容',
    '其他依赖：rsl_rl, tensorboard, wandb, gymnasium',
]
for item in sw_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('2.3 训练规模（论文Table VIII配置）', 2)
train_items = [
    '并行环境数：12,288个',
    'Episode步数：2,400步（24秒 @100Hz控制频率）',
    '每次迭代收集步数：12,288 x 24 = 294,912步',
    'PPO Epochs：2（每次迭代重复训练2轮）',
    '每次更新数据量：约5,900万状态-动作对',
    '学习率：1e-5（远小于常规PPO的3e-4）',
    'Batch Size：24步截断（per-env truncation window）',
    '预计训练总迭代：3000-5000 iterations',
]
for item in train_items:
    doc.add_paragraph(item, style='List Bullet')

# ==== Chapter 3: Limitations ====
doc.add_heading('三、原版复现存在的局限性与障碍', 1)

limitations = [
    ('3.1 最大障碍：Isaac Gym已不可获取',
     'Isaac Gym是NVIDIA的私有闭源软件，已于2024年停止维护更新。\n'
     'Preview 4版本需要通过NVIDIA开发者计划提交申请才能获取，普通用户无法直接下载。\n'
     '此外Isaac Gym仅支持Ubuntu Linux操作系统，无法在Windows/macOS上运行。\n'
     '这构成了原版复现最根本性的障碍——无法获取和使用核心仿真引擎。'),

    ('3.2 GPU资源需求远超当前设备能力',
     '论文使用12,288个并行环境同时在GPU上进行物理仿真，至少需要24GB以上显存（如A100 40GB/80GB或RTX 4090 24GB）。\n'
     '当前设备NVIDIA T600 Laptop仅有4GB VRAM，差距达6倍以上。\n'
     '即使大幅缩减环境数（如从12288降至512），训练效率和收敛性都会显著下降——\n'
     'PPO作为On-Policy算法，样本量直接影响优势估计的准确性。'),

    ('3.3 操作系统不兼容',
     '当前设备运行Windows 11。\n'
     'Isaac Gym = Linux only。若坚持原版方案，必须安装Ubuntu双系统或使用WSL2。\n'
     'WSL2虽然支持CUDA，但存在GPU直通性能损耗和Isaac Gym兼容性问题。'),

    ('3.4 机器人模型不匹配：XBot-S vs XBot-L',
     '论文使用XBot-S（1.2m身高/38kg），但humanoid-gym开源代码仅公开了XBot-L（1.65m/57kg）的URDF/MJCF模型。\n'
     '两者在质量分布、惯量矩阵、基座高度目标（0.7m vs 0.89m）、关节力矩范围等方面存在差异。\n'
     '论文中针对XBot-S调试的所有超参数（奖励权重、PD增益、域随机化范围）无法直接适用于XBot-L。'),

    ('3.5 闭式运动链的仿真近似问题',
     'XBot踝关节通过闭式运动链（四连杆机构parallel linkage）实现2DOF运动。\n'
     'Isaac Gym不支持闭式运动学约束，论文使用"虚拟电机"近似模拟。\n'
     '部署时通过"重映射(remap)"将虚拟电机目标转为实际电机指令。\n'
     '论文未详细说明该映射关系的数学形式（线性? 运动学解析? 数据拟合?），'
     '这是sim-to-real gap的关键来源之一。'),

    ('3.6 训练成本与可复现性',
     '12,288环境 x 2,400步 x 2 epoch = 约5,900万样本/迭代。\n'
     '完整训练需要3,000-5,000次迭代，总计约1,800亿-3,000亿样本。\n'
     '对个人或小型实验室而言，GPU集群的电费和硬件成本极高。\n'
     '论文未公开完整训练日志、随机种子和checkpoint文件，难以完全复现。'),

    ('3.7 仿真器耦合带来的迁移困难',
     'DWL的12项奖励函数中，部分指标直接依赖Isaac Gym的特定API：\n'
     '  - 足部接触力：Isaac Gym的contact_force tensor\n'
     '  - 地形高度扫描：Isaac Gym的heightfield API\n'
     '  - 域随机化参数：Isaac Gym的DR管理器\n'
     '换用其他物理引擎时，这些数值结果会有细微差异。虽然DWL的DR+瓶颈设计提供了鲁棒性，'
     '但奖励函数的微小数值偏差可能累积为不同的训练动力学。'),
]

for title, content in limitations:
    doc.add_heading(title, 2)
    doc.add_paragraph(content)

# ==== Chapter 4: Solution ====
doc.add_heading('四、解决方案：MuJoCo替代方案详解', 1)

doc.add_heading('4.1 方案核心思想', 2)
doc.add_paragraph(
    '保持DWL的算法完整性（网络架构/训练目标/奖励函数完全不变），'
    '仅将底层仿真引擎从Isaac Gym替换为MuJoCo（开源免费的工业级物理引擎）。\n\n'
    '这类似于"换引擎不换车身"——发动机（仿真器）换了，但驾驶方式（DWL算法）不变。')

doc.add_heading('4.2 完全保持不变的部分', 2)
keep_items = [
    'DWL网络架构：GRU Encoder(47->256->24) + Decoder(24->64->184) + Actor(24->48->12) + Critic(184->512->512->256->1)',
    'PPO算法：Clip Loss (epsilon=0.2) + GAE优势估计 (gamma=0.995, lambda=0.95)',
    'DWL联合训练目标：L_total = L_denoise + 5.0*L_pi + 5.0*L_v',
    'L_denoise = MSE(s_tilde, s_t) + 0.002 * ||z_t||_1（去噪重建+稀疏正则）',
    '五次多项式足部轨迹：f(t) = 9.6t^5 + 12t^4 - 18.8t^3 + 5t^2 + 0.1t',
    '12项奖励函数：线速度跟踪、角速度跟踪、姿态、身高、周期性接触力、足部高度跟踪等',
    '域随机化策略：摩擦系数U(0.2,2.0)、电机强度U(0.9,1.1)、载荷U(-5,20)kg、传感器噪声等',
    '双频控制架构：100Hz策略层(GRU+Actor) + 500Hz执行层(PD控制器)',
]
for item in keep_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.3 需要改变的部分', 2)
change_items = [
    '仿真引擎：Isaac Gym API → MuJoCo API（mj_step, data.qpos, data.qvel等）',
    '并行方式：GPU原生12288并行 → CPU多进程向量化并行（目标512-1024环境）',
    '机器人模型：XBot-S → XBot-L（使用humanoid-gym提供的MJCF文件）',
    '观测提取：gym.get_dof_state() → data.qpos[-12:] + 手动计算基底速度/欧拉角',
    'PD控制：Isaac Gym内置actuator → 手动实现 tau = Kp*(q_des - q) - Kd*qvel',
    '域随机化：gym.set_sim_params() → 创建仿真时修改model属性',
    '地形：Isaac Gym trimesh → MuJoCo hfield或自定义mesh',
]
for item in change_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.4 选择MuJoCo的具体理由', 2)
mujoco_reasons = [
    '免费且开源：DeepMind开发，Apache 2.0许可证，无需任何申请',
    '跨平台支持：Windows/Linux/macOS全支持，可在当前Win11设备直接运行',
    'MJCF模型已就绪：humanoid-gym仓库已提供XBot-L的MuJoCo MJCF模型文件和STL网格',
    '参考代码完备：humanoid-gym/scripts/sim2sim.py 展示了完整的MuJoCo推理循环（观测提取、PD控制、策略部署）',
    '物理精度高：工业级刚体动力学引擎，被广泛用于机器人学顶级研究',
    'GPU加速路径：MJX（MuJoCo XLA）可将物理计算编译到XLA在GPU上执行',
    '内存友好：单个MuJoCo实例内存占用远小于Isaac Gym，可在4GB显存设备上并行足够环境',
    '活跃维护：持续更新，社区活跃，文档完善',
]
for r in mujoco_reasons:
    doc.add_paragraph(r, style='List Bullet')

# ==== Chapter 5: Implementation Plan ====
doc.add_heading('五、具体实施步骤', 1)

steps_detail = [
    ('Step 1：从humanoid-gym提取模型文件',
     '执行git clone https://github.com/roboterax/humanoid-gym\n'
     '提取 resources/robots/XBot/mjcf/XBot-L.xml （MJCF模型）\n'
     '提取 resources/robots/XBot/meshes/ （全部STL网格文件）\n'
     '放置于 复现/models/ 目录下'),

    ('Step 2：安装CUDA + MuJoCo + GPU PyTorch',
     '安装 CUDA Toolkit 12.x（从NVIDIA官网下载）\n'
     '卸载CPU版PyTorch，安装GPU版：pip install torch torchvision --index-url https://download.pytorch.org/whl/cu121\n'
     '安装MuJoCo：pip install mujoco mujoco-viewer\n'
     '验证：python -c "import torch; print(torch.cuda.is_available()); import mujoco; print(mujoco.__version__)"'),

    ('Step 3：编写MuJoCo DWL环境类 (dwl_env.py)',
     '实现 XBotLMuJoCoEnv 类，包含：\n'
     '  - __init__()：加载MJCF模型，初始化MuJoCo data/model，设置PD参数\n'
     '  - reset()：重置机器人到初始姿态，随机化域参数（摩擦/质量/电机强度）\n'
     '  - step(action)：执行PD控制→推进物理→计算奖励→提取观测→判断终止\n'
     '  - _get_obs()：提取47维本体感知观测（论文Table I）\n'
     '  - _get_privileged_obs()：提取73维特权观测 + 拼接为完整状态\n'
     '  - _compute_rewards()：计算12项DWL奖励（论文Table V）\n'
     '  - _apply_domain_randomization()：应用DR参数'),

    ('Step 4：编写DWL网络模块 (dwl_networks.py)',
     '基于已有的dwl_reproduction.py架构，实现：\n'
     '  - GRUEncoder：GRU(input=47, hidden=256) → Linear(256,256) → ELU → Linear(256,24)\n'
     '  - Decoder：Linear(24,64) → ELU → Linear(64,184)\n'
     '  - Actor：Linear(24,48) → ELU → Linear(48,12) + 可学习log_std\n'
     '  - Critic：Linear(184,512) → ELU → Linear(512,512) → ELU → Linear(512,256) → ELU → Linear(256,1)'),

    ('Step 5：编写PPO+DWL训练器 (dwl_ppo.py)',
     '实现 DWLTrainer 类，包含：\n'
     '  - compute_gae()：广义优势估计\n'
     '  - update()：DWL联合训练（L_denoise + L_pi + L_v反向传播）\n'
     '  - 多环境并行rollout收集\n'
     '  - 学习率调度、梯度裁剪、checkpoint保存'),

    ('Step 6：编写主训练脚本 (train.py)',
     '实现完整的 main() 函数：\n'
     '  - 解析命令行参数（env数/学习率/迭代数等）\n'
     '  - 创建环境 + DWL模型 + PPO训练器\n'
     '  - 训练循环：rollout → GAE → DWL update → log → save\n'
     '  - TensorBoard / WandB日志记录\n'
     '  - 定期保存checkpoint\n'
     '  - 训练曲线可视化'),

    ('Step 7：冒烟测试',
     '验证完整流程可运行：\n'
     '  - 模型前向传播 + 反向传播正常\n'
     '  - MuJoCo环境reset/step正常\n'
     '  - 1个iteration完整训练不崩溃\n'
     '  - GPU内存使用在合理范围内\n'
     '  - 奖励值随训练有变化趋势（即使不收敛也算通过）'),
]

for title, content in steps_detail:
    doc.add_heading(title, 2)
    doc.add_paragraph(content)

# ==== Chapter 6: Device Adaptation ====
doc.add_heading('六、当前设备配置与适配方案', 1)

device_data = [
    ('GPU', 'NVIDIA T600 Laptop\n(4GB VRAM, 896 CUDA Cores)',
     '可跑512-1024并行环境\n网络训练用GPU\nMuJoCo物理在CPU算'),
    ('CPU', 'AMD Ryzen（含集成显卡）', '多核跑MuJoCo物理步进\n可用多进程并行'),
    ('操作系统', 'Windows 11 Home China', 'MuJoCo原生支持Windows\n无需装双系统/WSL'),
    ('PyTorch', '2.11.0 CPU版（当前）', '需重装为CUDA 12.1 GPU版\npip install torch --index-url cu121'),
    ('CUDA', '未安装Toolkit\nnvcc不可用', '安装CUDA 12.1 Toolkit\nNVIDIA官网下载'),
    ('仿真器', '无Isaac Gym\n有Gymnasium 1.3.0', 'pip install mujoco>=3.0\nmjcf模型已就绪'),
]
table = doc.add_table(rows=len(device_data) + 1, cols=3, style='Light Grid Accent 1')
table.rows[0].cells[0].text = '项目'
table.rows[0].cells[1].text = '当前状态'
table.rows[0].cells[2].text = '适配方案'
for i, (item, status, fix) in enumerate(device_data):
    table.rows[i + 1].cells[0].text = item
    table.rows[i + 1].cells[1].text = status
    table.rows[i + 1].cells[2].text = fix

doc.add_heading('6.1 显存预算估算', 2)
est_data = [
    ('DWL模型参数', '~320K', '~1.3MB (fp32)', 'GPU'),
    ('优化器状态(Adam)', '~320K x 2', '~2.6MB (fp32)', 'GPU'),
    ('Rollout Buffer(512 env × 24 steps)', '512×24×231维', '~11MB (fp32)', 'GPU'),
    ('MuJoCo物理状态(512 env)', '~512×2000维', '~4MB (fp32)', 'CPU'),
    ('总计GPU显存', '-', '~15MB (模型+Buffer)', '-'),
]
table = doc.add_table(rows=len(est_data) + 1, cols=4, style='Light Grid Accent 1')
table.rows[0].cells[0].text = '组件'
table.rows[0].cells[1].text = '规模'
table.rows[0].cells[2].text = '占用'
table.rows[0].cells[3].text = '位置'
for i, (comp, scale, mem, loc) in enumerate(est_data):
    table.rows[i + 1].cells[0].text = comp
    table.rows[i + 1].cells[1].text = scale
    table.rows[i + 1].cells[2].text = mem
    table.rows[i + 1].cells[3].text = loc
doc.add_paragraph('结论：4GB VRAM完全足够。512环境的GPU占用不到100MB，远在T600的能力范围内。瓶颈在CPU物理计算而非GPU。')

# ==== Chapter 7: File Structure ====
doc.add_heading('七、复现代码文件结构', 1)
doc.add_paragraph('Humanoid Locomotion/复现/ 目录结构：')
code = (
    '复现/\n'
    '├── models/\n'
    '│   ├── XBot-L.xml              # MuJoCo MJCF模型（从humanoid-gym复制）\n'
    '│   └── meshes/                 # STL网格文件（80个文件）\n'
    '├── dwl_env.py                  # MuJoCo环境类（~400行）\n'
    '│   ├── XBotLMuJoCoEnv          # reset/step/观测提取/奖励计算/DR\n'
    '│   └── VecEnv                  # 多进程并行环境包装器\n'
    '├── dwl_networks.py             # DWL网络架构（~200行）\n'
    '│   ├── GRUEncoder              # GRU(47→256) + MLP(256→24)\n'
    '│   ├── Decoder                 # MLP(24→64→184)\n'
    '│   ├── Actor                   # MLP(24→48→12) + log_std\n'
    '│   └── Critic                  # MLP(184→512→512→256→1)\n'
    '├── dwl_ppo.py                  # PPO+DWL训练器（~300行）\n'
    '│   ├── compute_gae()           # GAE优势估计\n'
    '│   └── DWLTrainer.update()     # L_denoise + L_pi + L_v联合更新\n'
    '├── dwl_config.py               # 超参数配置（~80行）\n'
    '│   └── DWLConfig               # 论文Table II/V/VIII参数\n'
    '├── train.py                    # 主训练脚本（~150行）\n'
    '│   └── main()                  # 训练主循环\n'
    '├── dwl_reproduction.py         # 原有架构验证代码（保留）\n'
    '├── logs/                       # TensorBoard日志\n'
    '├── checkpoints/                # 模型保存\n'
    '└── requirements.txt            # Python依赖列表\n'
)
doc.add_paragraph(code)

# ==== Chapter 8: Key Architecture ====
doc.add_heading('八、DWL核心架构速览', 1)
doc.add_paragraph(
    'DWL与前身算法的演进路径：\n'
    '  A3C(异步优势Actor-Critic, 2016) → A2C(同步版, 2016) → PPO(+Clipping, 2017) → DWL(+Denoising, 2024)\n\n'
    '重要澄清：\n'
    '  A3C的"A" = Asynchronous（异步并行训练）\n'
    '  DWL的"Asymmetric Actor-Critic"中的"A" = Asymmetric（非对称：Actor看47维观测，Critic看184维完整状态）\n'
    '  两者是完全不同的概念！\n\n'
    'DWL核心公式：\n'
    '  L_total = L_denoise + 5.0 * L_pi + 5.0 * L_v\n'
    '  L_denoise = ||s_tilde - s_t||² + 0.002 * ||z_t||₁\n'
    '  L_pi = -min(r_t * A_t, clip(r_t, 0.8, 1.2) * A_t)\n'
    '  L_v = MSE(V(s_t), R_t)\n'
    '  r_t = pi_new(a_t|z_t) / pi_old(a_t|z_t)\n\n'
    '隐状态瓶颈：\n'
    '  47维观测 → GRU(256维) → 压缩 → z_t(24维隐状态)\n'
    '  z_t同时喂给Decoder(重建184维状态)、Actor(输出12维动作)、Critic(估计价值)\n\n'
    '部署时仅需Encoder + Actor：\n'
    '  47维观测 → GRU → z_t(24维) → Actor → 12维关节目标位置 → PD控制器(500Hz)'
)

# ==== Chapter 9: Summary ====
doc.add_heading('九、总结', 1)
doc.add_paragraph(
    'DWL（Denoising World Model Learning）是2024年RSS会议上发表的里程碑式工作，'
    '首次实现了人形机器人在复杂地形（雪地、楼梯、不规则路面）上的零样本Sim-to-Real迁移。\n\n'
    '原版复现面临的最大障碍是Isaac Gym的获取问题和极高的硬件门槛（12288并行环境，24GB+显存）。\n\n'
    '本方案通过在MuJoCo开源物理引擎上重建DWL训练流程，在保持算法完整性（网络架构、训练目标、'
    '奖励函数完全不变）的前提下，使复现能够在普通硬件（NVIDIA T600 4GB / Windows 11）上进行。\n\n'
    '主要权衡：\n'
    '  1. 环境并行数从12,288降至512-1024，训练效率降低但可运行\n'
    '  2. 机器人模型从XBot-S变为XBot-L（开源模型限制），超参数需微调\n'
    '  3. 物理引擎从Isaac Gym换为MuJoCo，仿真动力学有细微差异但DWL的鲁棒设计可吸收\n\n'
    '核心优势：\n'
    '  1. 完全避开Isaac Gym的获取和平台限制，可在Windows上运行\n'
    '  2. MuJoCo免费开源，长期可用，社区活跃\n'
    '  3. humanoid-gym官方代码已提供MJCF模型和MuJoCo推理参考\n'
    '  4. 算法层面与论文完全一致，学术复现价值不打折\n'
    '  5. 4GB显存即可训练，大幅降低硬件门槛'
)

output_path = desktop + r'\Humanoid Locomotion\DWL复现方案总结.docx'
doc.save(output_path)
print(f'Word document saved to: {output_path}')
